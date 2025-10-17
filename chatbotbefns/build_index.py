import os
import re
import unicodedata
import sys
from pathlib import Path
import docx
from dotenv import load_dotenv
from tqdm import tqdm
from datetime import datetime, timezone

# --- Vector DB / Embeddings ---
from pinecone import Pinecone, ServerlessSpec
from openai import OpenAI

# =========================
# Config
# =========================
INDEX_NAME = "gh-index"
EMBED_MODEL = "text-embedding-3-small"  # 1536-dim; switch to "-large" (3072) if you prefer
MODEL_DIMS = {"text-embedding-3-small": 1536, "text-embedding-3-large": 3072}
NAMESPACE = os.getenv("PC_NAMESPACE", "")  # optional; leave "" to use default

# Source docs (make these absolute or relative to this script)
SPECIFIC_DOCX_FILES = [
    "Cottage_Specific_Information.docx",
]
STANDARD_DOCX = "Information_Standard_Cottages.docx"

# =========================
# Init
# =========================
load_dotenv()

def die(msg: str):
    print(f"ERROR: {msg}", file=sys.stderr)
    sys.exit(1)

# Resolve paths relative to the script’s directory (safer than CWD)
BASE_DIR = Path(__file__).resolve().parent
SPECIFIC_DOCX_FILES = [str((BASE_DIR / p).resolve()) for p in SPECIFIC_DOCX_FILES]
STANDARD_DOCX = str((BASE_DIR / STANDARD_DOCX).resolve())

# Check files exist early
missing = [p for p in SPECIFIC_DOCX_FILES + [STANDARD_DOCX] if not Path(p).exists()]
if missing:
    die(f"Missing input .docx file(s): {missing}")

openai_key = os.getenv("OPENAI_API_KEY")
pc_key = os.getenv("PINECONE_API_KEY")
if not openai_key:
    die("OPENAI_API_KEY not set.")
if not pc_key:
    die("PINECONE_API_KEY not set.")

client = OpenAI(api_key=openai_key)
pc = Pinecone(api_key=pc_key)
pinecone_region = os.getenv("PINECONE_ENV", "us-east-1")
index_dim = MODEL_DIMS[EMBED_MODEL]

# Create Pinecone index if needed, then wait until ready
existing = pc.list_indexes().names()
if INDEX_NAME not in existing:
    print(f"Creating Pinecone index '{INDEX_NAME}' (dim={index_dim}) in {pinecone_region}...")
    pc.create_index(
        name=INDEX_NAME,
        dimension=index_dim,
        metric="cosine",
        spec=ServerlessSpec(cloud="aws", region=pinecone_region),
    )

# Wait for readiness (important with serverless)
print("Waiting for index to be ready...")
import time
while True:
    desc = pc.describe_index(INDEX_NAME)
    status = getattr(desc, "status", None)
    ready = False
    if status is not None:
        ready = getattr(status, "ready", False)
        try:
            ready = ready or bool(status.get("ready", False))  # type: ignore[attr-defined]
        except Exception:
            pass
    if ready:
        break
    time.sleep(1)

index = pc.Index(INDEX_NAME)

# =========================
# Utilities
# =========================
def slugify(value: str) -> str:
    value = unicodedata.normalize("NFKD", value)
    value = re.sub(r"[^\w\s-]", "", value, flags=re.UNICODE).strip().lower()
    value = re.sub(r"[-\s]+", "-", value)
    return value

def normalize_title(t: str) -> str:
    t = t.strip().rstrip(":")
    # remove trailing 'Specific'
    t = re.sub(r"\s+Specific$", "", t, flags=re.IGNORECASE)
    # fix known misspells/variants
    t = re.sub(r"\btrealce\b", "Treacle", t, flags=re.IGNORECASE)
    # collapse whitespace
    t = re.sub(r"\s+", " ", t)
    # Title Case
    t = t.title()
    # re-upcase common acronyms if needed
    t = re.sub(r"\bTv\b", "TV", t)
    return t

# Fallback title detector to catch headings like "Lanty Scar" (no suffix) and "4 Studio House"
TITLE_FALLBACK = re.compile(
    r"""^\s*
        (?:\d+\s+)?                                 # optional leading number (e.g., "4 ")
        [A-Z][A-Za-z']+(?:\s+[A-Z][A-Za-z']+){0,3}  # 1–4 Title-Case words
        (?:\s+(?:Cottage|House|Flat|View|Mews|Fold))?  # optional suffix
        \s*$
    """,
    re.IGNORECASE | re.VERBOSE
)

GLOBAL_NOISE_HEADINGS = {
    "FIRE ALARM TIMES", "FIRE ALARM TIMES AGREED WITH EMCS", "WHEELCHAIR HIRE"
}

NUM_PREFIX = re.compile(r"^\s*\d+[\.\)\-–:]?\s*")  # e.g., "1.", "2)", "3 –"

def clean_heading_text(txt: str) -> str:
    # remove auto-numbering / bullets at the start of headings
    return NUM_PREFIX.sub("", txt or "").strip()

def is_valid_heading_text(txt: str) -> bool:
    # must contain at least one letter (avoid "1", "2.", etc.)
    return bool(re.search(r"[A-Za-z]", txt))

def read_paragraphs(file_path: str):
    doc = docx.Document(file_path)
    return doc.paragraphs

def extract_cottages_from_docx(file_path: str):
    """
    Prefer Word heading styles; fallback to regex-only headings.
    Merge duplicate sections by normalized title.
    Ignores numeric-only headings and strips list numbering.
    """
    paragraphs = read_paragraphs(file_path)
    by_slug = {}
    current_slug = None
    buffer = []

    def flush():
        nonlocal current_slug, buffer
        if current_slug and buffer:
            chunk = "\n".join([ln for ln in buffer if ln.strip()]).strip()
            if chunk:
                by_slug[current_slug]["text"] = (by_slug[current_slug]["text"] + "\n" + chunk).strip()
        buffer = []

    for p in paragraphs:
        raw = (p.text or "").strip()
        if not raw:
            continue

        # detect heading by style OR by fallback regex
        try:
            style_name = (p.style.name or "").lower()
        except Exception:
            style_name = ""
        is_heading_style = style_name.startswith("heading")

        cleaned = clean_heading_text(raw)
        looks_like_title = bool(TITLE_FALLBACK.match(cleaned))

        if is_heading_style or looks_like_title:
            # skip global/non-cottage headings
            if any(h == raw.upper() for h in GLOBAL_NOISE_HEADINGS):
                continue
            # skip headings that are just numbers / no letters
            if not is_valid_heading_text(cleaned):
                continue

            # new section
            flush()
            title = normalize_title(cleaned)
            slug = slugify(title)
            if slug not in by_slug:
                by_slug[slug] = {"title": title, "text": ""}
            current_slug = slug
        else:
            buffer.append(raw)

    flush()
    # drop empties
    return {s: v for s, v in by_slug.items() if v["text"].strip()}

def read_docx_as_text(file_path: str):
    doc = docx.Document(file_path)
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    return "\n".join(lines)

def format_standard_info(raw_text: str) -> str:
    raw_text = raw_text.replace("\r", "").strip()
    lines = [ln.strip() for ln in raw_text.split("\n") if ln.strip()]
    out = ["#### 🏠 General Information for All Cottages", ""]
    for ln in lines:
        out.append(f"- {ln}")
    return "\n".join(out)

def parse_fields(text: str):
    """
    Parse "Key: Value" style lines; collect others in 'extra'.
    Accepts bullets/dashes and en-dash.
    """
    fields = {}
    extra = []
    for raw in [l.strip("-•").strip() for l in text.split("\n") if l.strip()]:
        m = re.match(r"^([A-Za-z][A-Za-z0-9\s/()&]+?)\s*[:\-–]\s*(.+)$", raw)
        if m:
            key = m.group(1).strip().lower().replace(" ", "_")
            val = m.group(2).strip()
            fields[key] = val
        else:
            extra.append(raw)
    return fields, extra

def generate_embedding_text(title, fields, extra_lines):
    lines = [f"Cottage: {title}"]
    for k, v in fields.items():
        pretty = k.replace("_", " ").capitalize()
        lines.append(f"{pretty}: {v}")
    if extra_lines:
        lines.extend(extra_lines)
    return "\n".join(lines)

def get_embedding(text: str):
    resp = client.embeddings.create(
        input=text,
        model=EMBED_MODEL
    )
    emb = resp.data[0].embedding
    if not isinstance(emb, list) or not emb:
        raise RuntimeError("Empty embedding received")
    return emb

# =========================
# Load & merge specific info (supports multiple files)
# =========================
specific_map = {}  # slug -> {"title": str, "text": str}
for fp in SPECIFIC_DOCX_FILES:
    sections = extract_cottages_from_docx(fp)  # dict[slug] = {"title","text"}
    for slug, payload in sections.items():
        if slug not in specific_map:
            specific_map[slug] = payload
        else:
            # merge additional body text if new
            existing = specific_map[slug]["text"]
            new_text = payload["text"].strip()
            if new_text and new_text not in existing:
                specific_map[slug]["text"] = (existing + "\n" + new_text).strip()

standard_text_raw = read_docx_as_text(STANDARD_DOCX)
standard_text_formatted = format_standard_info(standard_text_raw)

print(f"Parsed cottages: {len(specific_map)}")
for i, (slug, v) in enumerate(list(specific_map.items())[:20]):
    print(f"  - {i+1}. {v['title']} (slug={slug})")

if len(specific_map) == 0:
    die("No cottage sections parsed. Check heading detection or source documents.")

# =========================
# Build vectors & upsert
# =========================
ups = []
print("Generating embeddings and preparing vectors...")
for slug, item in tqdm(list(specific_map.items()), total=len(specific_map)):
    title = item["title"]
    body = item["text"]

    fields, extra_lines = parse_fields(body)
    embedding_text = generate_embedding_text(title, fields, extra_lines)
    emb = get_embedding(embedding_text)

    # sanity check embedding length matches index_dim
    if len(emb) != index_dim:
        die(f"Embedding dim {len(emb)} != index dim {index_dim} for '{title}'")

    metadata = {
        "title": title,
        "text": embedding_text,
        "standard_info": standard_text_formatted,
        "InformationSourceDate": datetime.now(timezone.utc).isoformat(),
        **fields
    }

    uid = f"cottage-{slug}"  # stable across runs
    ups.append({
        "id": uid,
        "values": emb,
        "metadata": metadata,
    })

print(f"Prepared {len(ups)} vectors to upsert.")
if not ups:
    die("Nothing to upsert (no vectors prepared).")

# Chunked upserts (safer for larger batches)
BATCH = 100
for i in range(0, len(ups), BATCH):
    batch = ups[i:i+BATCH]
    index.upsert(vectors=batch, namespace=NAMESPACE if NAMESPACE else None)

# Confirm with stats
stats = index.describe_index_stats(namespace=NAMESPACE if NAMESPACE else None)
print("Index stats:", stats)

# Optional: quick query sanity check (comment out if not desired)
try:
    probe_vec = ups[0]["values"]
    probe = index.query(
        vector=probe_vec,
        top_k=3,
        include_metadata=True,
        namespace=NAMESPACE if NAMESPACE else None
    )
    top_ids = [m.get("id") if isinstance(m, dict) else getattr(m, "id", None) for m in probe.get("matches", [])]
    print("Sample query OK. Top ids:", top_ids)
except Exception as e:
    print("Sample query failed (non-fatal):", e)

print("✅ Indexing complete with numeric-heading cleanup, robust title detection, slug dedupe, and updated embeddings.")
